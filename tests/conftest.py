"""Pytest configuration and session-scoped fixtures.

Binary test files (PDF, DOCX) are generated programmatically the first time
the test suite runs and cached on disk.  Delete tests/files/sample.pdf or
tests/files/sample.docx to force regeneration.
"""
from __future__ import annotations

import sys
from pathlib import Path

# Ensure the project root is on sys.path so `app` can be imported without
# an editable install.
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest

FILES_DIR = Path(__file__).parent / "files"

# ---------------------------------------------------------------------------
# PDF generator
# ---------------------------------------------------------------------------

def _create_pdf(path: Path) -> None:
    if path.exists():
        return

    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    cur_y: float = 60.0
    left: float = 60.0

    def put(t: str, fontsize: float = 11.0, fontname: str = "helv") -> None:
        nonlocal cur_y
        page.insert_text((left, cur_y), t, fontsize=fontsize, fontname=fontname)
        cur_y += fontsize * 1.5

    def skip(pts: float = 15.0) -> None:
        nonlocal cur_y
        cur_y += pts

    # H1 (fontsize 22 → heading level 1 when body is 11)
    put("Python Web Frameworks", fontsize=22)
    skip()

    # H2 Introduction  (fontsize 16 → heading level 2)
    put("Introduction", fontsize=16)
    put("Web frameworks are software libraries designed to support the development")
    put("of web applications, including web services, web resources, and web APIs.")
    put("They handle common tasks such as routing, templating, and authentication.")
    put("Choosing a framework early in a project has long-term architectural impact.")
    skip()

    # H2 Popular Frameworks
    put("Popular Frameworks", fontsize=16)

    # H3 Flask  (fontsize 14 → heading level 3)
    put("Flask", fontsize=14)
    put("Flask is a lightweight WSGI web application framework designed to make")
    put("getting started quick and easy, scaling up to complex applications.")
    skip()

    # H3 Django
    put("Django", fontsize=14)
    put("Django is a high-level Python web framework that encourages rapid")
    put("development and clean, pragmatic design built by experienced developers.")
    skip()

    # H2 Comparison + drawn table
    put("Comparison", fontsize=16)

    table_data = [
        ["Framework", "Type",  "First Release"],
        ["Flask",     "Micro", "2010"],
        ["Django",    "Full",  "2005"],
        ["FastAPI",   "Async", "2018"],
    ]
    col_w, row_h = 130, 18
    cols_n, rows_n = 3, len(table_data)
    t_x0, t_y0 = left, cur_y

    # Draw table grid (detected by page.find_tables() in PyMuPDF ≥ 1.23)
    shape = page.new_shape()
    for r in range(rows_n + 1):
        yr = t_y0 + r * row_h
        shape.draw_line((t_x0, yr), (t_x0 + cols_n * col_w, yr))
    for c in range(cols_n + 1):
        xc = t_x0 + c * col_w
        shape.draw_line((xc, t_y0), (xc, t_y0 + rows_n * row_h))
    shape.finish(color=(0, 0, 0), width=0.5)
    shape.commit()

    # Table cell text at fontsize=10 (distinct from body=11 and code=8)
    for r, row in enumerate(table_data):
        for c, cell in enumerate(row):
            page.insert_text(
                (t_x0 + c * col_w + 4, t_y0 + r * row_h + 13),
                cell,
                fontsize=10,
            )
    cur_y = t_y0 + rows_n * row_h
    skip(30)  # extra gap so PyMuPDF places the next heading in its own text block

    # H2 Key Features + list items
    put("Key Features", fontsize=16)
    for feat in [
        "- URL routing and request dispatching",
        "- Template rendering engine",
        "- Database abstraction layer",
        "- Built-in authentication and authorization",
    ]:
        put(feat)
    skip()

    # H2 Code Example + code text
    put("Code Example", fontsize=16)
    for line in [
        "from flask import Flask",
        "app = Flask(__name__)",
        "",
        '@app.route("/")',
        "def hello():",
        '    return "Hello, World!"',
    ]:
        put(line or " ", fontsize=8, fontname="cour")
    skip()

    # H2 Conclusion + paragraph
    put("Conclusion", fontsize=16)
    put("Choosing the right web framework depends on your project requirements")
    put("and team expertise. Each framework has its own philosophy and trade-offs.")

    doc.save(str(path))
    doc.close()


# ---------------------------------------------------------------------------
# DOCX generator
# ---------------------------------------------------------------------------

def _create_docx(path: Path) -> None:
    if path.exists():
        return

    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt

    doc = Document()

    # Add a "Code" paragraph style when the default template does not include it
    try:
        doc.styles["Code"]
    except KeyError:
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(9)

    doc.add_heading("Python Web Frameworks", level=1)

    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "Web frameworks are software libraries designed to support the development "
        "of web applications, including web services, web resources, and web APIs. "
        "They handle common tasks such as routing, templating, and authentication."
    )

    doc.add_heading("Popular Frameworks", level=2)

    doc.add_heading("Flask", level=3)
    doc.add_paragraph(
        "Flask is a lightweight WSGI web application framework. It is designed to "
        "make getting started quick and easy, with the ability to scale up to "
        "complex applications. Flask is classified as a micro-framework."
    )

    doc.add_heading("Django", level=3)
    doc.add_paragraph(
        "Django is a high-level Python web framework that encourages rapid "
        "development and clean, pragmatic design. Built by experienced developers, "
        "it takes care of much of the hassle of web development."
    )

    doc.add_heading("Comparison", level=2)

    table = doc.add_table(rows=4, cols=3)
    headers = ["Framework", "Type", "First Release"]
    rows_data = [
        ["Flask",   "Micro", "2010"],
        ["Django",  "Full",  "2005"],
        ["FastAPI", "Async", "2018"],
    ]
    for c, header in enumerate(headers):
        table.rows[0].cells[c].text = header
    for r, row_data in enumerate(rows_data, start=1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_heading("Key Features", level=2)
    for feature in [
        "URL routing and request dispatching",
        "Template rendering engine",
        "Database abstraction layer",
        "Built-in authentication and authorization",
    ]:
        doc.add_paragraph(feature, style="List Bullet")

    doc.add_heading("Code Example", level=2)

    # Consecutive "Code" paragraphs are merged into one CODE block by the parser
    for line in [
        "from flask import Flask",
        "app = Flask(__name__)",
        " ",
        '@app.route("/")',
        "def hello():",
        '    return "Hello, World!"',
    ]:
        doc.add_paragraph(line, style="Code")

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "Choosing the right web framework depends on your project requirements, "
        "team expertise, and scalability needs. Flask offers simplicity while "
        "Django provides a full-featured ecosystem out of the box."
    )

    doc.save(str(path))


# ---------------------------------------------------------------------------
# Session fixture — generate binary files once
# ---------------------------------------------------------------------------

@pytest.fixture(scope="session", autouse=True)
def _generated_files() -> None:
    FILES_DIR.mkdir(parents=True, exist_ok=True)
    _create_pdf(FILES_DIR / "sample.pdf")
    _create_docx(FILES_DIR / "sample.docx")


# ---------------------------------------------------------------------------
# Path fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(scope="session")
def sample_txt_path() -> Path:
    return FILES_DIR / "sample.txt"


@pytest.fixture(scope="session")
def plain_txt_path() -> Path:
    return FILES_DIR / "plain.txt"


@pytest.fixture(scope="session")
def sample_html_path() -> Path:
    return FILES_DIR / "sample.html"


@pytest.fixture(scope="session")
def sample_pdf_path(_generated_files) -> Path:
    return FILES_DIR / "sample.pdf"


@pytest.fixture(scope="session")
def sample_docx_path(_generated_files) -> Path:
    return FILES_DIR / "sample.docx"
