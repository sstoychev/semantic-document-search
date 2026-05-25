"""DOCX parser — uses python-docx.

Extraction strategy:
  - Iterate doc.element.body children in document order to interleave
    paragraphs and tables correctly (doc.paragraphs and doc.tables come in
    separate flat lists that lose relative ordering).
  - Paragraph classification (by style name):
      "Heading N"            → HEADING level N
      "List Bullet" / "List Number" /
        "List Paragraph" or any style
        starting with "List", or paragraph
        has <w:numPr>          → LIST  (consecutive items merged into one block)
      style contains "code" /
        "preformatted" / "mono" → CODE  (consecutive lines merged into one block)
      everything else          → PARAGRAPH
  - Tables are converted to plain text: cells separated by " | ", rows by "\\n".
  - Heading hierarchy is tracked to build breadcrumbs.

Metadata: {} (no extras for DOCX beyond what the Block already carries)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.chunker.models import Block, BlockType


# ── Helpers ───────────────────────────────────────────────────────────────────

def _table_to_text(table: Table) -> str:
    rows: list[str] = []
    seen: set[str] = set()  # deduplicate merged-cell repetitions
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        row_key = "\t".join(cells)
        if row_key in seen:
            continue
        seen.add(row_key)
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _has_numpr(para: Paragraph) -> bool:
    """Return True if the paragraph has list numbering markup.

    <w:numPr> is a grandchild of <w:p> (lives inside <w:pPr>), so we use
    an XPath descendant search (.//…) rather than a direct-child find().
    """
    return para._p.find(f".//{qn('w:numPr')}") is not None


def _update_headings(heading_levels: dict[int, str], level: int, text: str) -> list[str]:
    heading_levels[level] = text
    for lvl in [lv for lv in heading_levels if lv > level]:
        del heading_levels[lvl]
    return [heading_levels[lv] for lv in sorted(heading_levels)]


_CODE_STYLE_KEYWORDS = ("code", "preformatted", "mono", "verbatim")


def _is_code_style(style_name: str) -> bool:
    lower = style_name.lower()
    return any(kw in lower for kw in _CODE_STYLE_KEYWORDS)


# ── Public entry point ────────────────────────────────────────────────────────

def parse_docx(path: Path) -> list[Block]:
    doc = Document(str(path))
    blocks: list[Block] = []
    position = 0
    heading_levels: dict[int, str] = {}

    def crumbs() -> list[str]:
        return [heading_levels[lv] for lv in sorted(heading_levels)]

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        # ── Paragraph ─────────────────────────────────────────────────────
        if tag == "p":
            para = Paragraph(element, doc)
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Heading
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1
                all_crumbs = _update_headings(heading_levels, level, text)
                blocks.append(Block(
                    position=position,
                    type=BlockType.HEADING,
                    breadcrumbs=all_crumbs[:-1],
                    text=text,
                    level=level,
                ))
                position += 1

            # Code
            elif _is_code_style(style_name):
                # Merge consecutive code paragraphs into a single block
                if blocks and blocks[-1].type == BlockType.CODE and blocks[-1].breadcrumbs == crumbs():
                    blocks[-1].text += f"\n{text}"
                else:
                    blocks.append(Block(
                        position=position,
                        type=BlockType.CODE,
                        breadcrumbs=crumbs(),
                        text=text,
                    ))
                    position += 1

            # List — covers "List Bullet", "List Number", "List Paragraph", etc.
            elif _has_numpr(para) or style_name.startswith("List"):
                # Merge consecutive list items into one LIST block
                if blocks and blocks[-1].type == BlockType.LIST and blocks[-1].breadcrumbs == crumbs():
                    blocks[-1].text += f"\n- {text}"
                else:
                    blocks.append(Block(
                        position=position,
                        type=BlockType.LIST,
                        breadcrumbs=crumbs(),
                        text=f"- {text}",
                    ))
                    position += 1

            # Regular paragraph
            else:
                blocks.append(Block(
                    position=position,
                    type=BlockType.PARAGRAPH,
                    breadcrumbs=crumbs(),
                    text=text,
                ))
                position += 1

        # ── Table ─────────────────────────────────────────────────────────
        elif tag == "tbl":
            table = Table(element, doc)
            text = _table_to_text(table)
            if text.strip():
                blocks.append(Block(
                    position=position,
                    type=BlockType.TABLE,
                    breadcrumbs=crumbs(),
                    text=text,
                ))
                position += 1

    return blocks
